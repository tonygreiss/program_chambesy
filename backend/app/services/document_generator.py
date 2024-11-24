from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from convertdate import coptic
import io
from datetime import datetime
import os
import csv

class DocumentGenerator:
    def __init__(self):
        """Initialize the document generator"""
        self.doc = Document()
        self._setup_document()
        
    def _setup_document(self):
        """Configure initial document settings"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def _get_coptic_date(self, year: int, month: int, day: int) -> str:
        """Convert Gregorian date to Coptic date string"""
        coptic_date = coptic.from_gregorian(year, month, day)
        coptic_month_names = {
            1: "Tout", 2: "Babah", 3: "Hatour", 4: "Kiahk",
            5: "Toubah", 6: "Amshir", 7: "Baramhat", 8: "Baramoudah",
            9: "Bashans", 10: "Paoni", 11: "Epip", 12: "Mesra", 13: "Nasie"
        }
        coptic_year = coptic_date[0]
        return f"{coptic_date[2]} {coptic_month_names[coptic_date[1]]} {coptic_year}"

    def _get_french_weekday(self, date_obj: datetime) -> str:
        """Get French weekday name"""
        french_weekdays = {
            0: "Lundi",
            1: "Mardi",
            2: "Mercredi",
            3: "Jeudi",
            4: "Vendredi",
            5: "Samedi",
            6: "Dimanche"
        }
        return french_weekdays[date_obj.weekday()]

    def _get_formatted_gregorian_date(self, year: int, month: int, day: int) -> str:
        """Format Gregorian date with French weekday"""
        date_obj = datetime(year, month, day)
        weekday = self._get_french_weekday(date_obj)
        return f"{weekday} {day:02d}-{month:02d}"

    def _set_cell_background(self, cell, color):
        """Set cell background color"""
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _set_table_style(self, table):
        """Apply custom table styling"""
        # Set table width to 100% of page width
        table.autofit = False
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(6.5 / len(row.cells))  # Distribute width evenly
        
        # Set table borders
        table.style = 'Table Grid'
        
        # Create a custom style for the table
        table._element.xpath('./w:tblPr')[0].append(
            parse_xml('''
                <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                    <w:left w:val="none"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                    <w:right w:val="none"/>
                    <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                    <w:insideV w:val="none"/>
                </w:tblBorders>
            ''')
        )

    def _load_weekly_schedule(self):
        """Load weekly church schedule from CSV"""
        schedule_dict = {}
        
        current_dir = os.path.dirname(os.path.abspath(__file__))
        csv_path = os.path.join(current_dir, '..', 'input files', 'church_schedule.csv')
        
        try:
            with open(csv_path, 'r', encoding='utf-8-sig') as file:
                reader = csv.DictReader(file)
                for row in reader:
                    weekday_num = int(row['week_day_num'])
                    schedule_dict[weekday_num] = {
                        'event': row['event'].strip() if row['event'] else '',
                        'time': row['time'].strip() if row['time'] else ''
                    }
                print(f"Loaded schedule: {schedule_dict}")
        except Exception as e:
            print(f"Error loading church schedule: {str(e)}")
            import traceback
            traceback.print_exc()
            return {}
        
        return schedule_dict

    def generate(self, year: int, month: int, dates: list, french_verse: str, arabic_verse: str) -> bytes:
        """Generate a church program document"""
        # Load weekly schedule
        weekly_schedule = self._load_weekly_schedule()
        print(f"Weekly schedule loaded: {weekly_schedule}")
        
        # Add title and verses (unchanged)
        title = self.doc.add_heading(f'Programme du mois {month}/{year}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        verse_para = self.doc.add_paragraph()
        verse_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        verse_para.add_run(french_verse).bold = True
        verse_para.add_run('\n')
        verse_para.add_run(arabic_verse).bold = True
        
        self.doc.add_paragraph()
        
        # Create and style table
        table = self.doc.add_table(rows=1, cols=4)
        self._set_table_style(table)
        
        # Set header row
        header_cells = table.rows[0].cells
        headers = ['Date', 'Heure', 'Evènement', 'Synaxaire']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Get days in month
        if month == 2:
            days_in_month = 29 if year % 4 == 0 and (year % 100 != 0 or year % 400 == 0) else 28
        elif month in [4, 6, 9, 11]:
            days_in_month = 30
        else:
            days_in_month = 31
        
        # Create synaxaire lookup dictionary
        synaxaire_dict = {int(entry['date'].split('/')[0]): entry for entry in dates}
        
        # Add data rows with alternating colors
        for day in range(1, days_in_month + 1):
            row_cells = table.add_row().cells
            
            # Set alternating row background
            if day % 2 == 0:
                for cell in row_cells:
                    self._set_cell_background(cell, "F2F2F2")
            
            # Get the date object to determine the day of week
            date_obj = datetime(year, month, day)
            weekday_num = date_obj.weekday() + 1  # Monday = 1, Sunday = 7
            print(f"Processing {date_obj.strftime('%A')}, weekday_num: {weekday_num}")
            
            # Date column
            gregorian_date = self._get_formatted_gregorian_date(year, month, day)
            coptic_date = self._get_coptic_date(year, month, day)
            row_cells[0].text = f"{gregorian_date}\n{coptic_date}"
            
            # Hour and Event columns from weekly schedule
            if weekday_num in weekly_schedule:
                schedule = weekly_schedule[weekday_num]
                print(f"Found schedule for {date_obj.strftime('%A')}: {schedule}")
                
                # Format multi-line event text
                event_text = schedule['event'].replace('\n', '\n') if schedule['event'] else ''
                
                row_cells[1].text = schedule['time']
                row_cells[2].text = event_text
                
                # Ensure proper formatting for multi-line text
                for cell in [row_cells[1], row_cells[2]]:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                row_cells[1].text = ""
                row_cells[2].text = ""
            
            # Synaxaire column
            if day in synaxaire_dict:
                first_entry = synaxaire_dict[day]['event'].split('\n')[0]
                if ': ' in first_entry:
                    row_cells[3].text = first_entry.split(': ', 1)[1]
                else:
                    row_cells[3].text = first_entry
            
            # Center align all cells and adjust line breaks
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Replace any '\n' in the text with proper line breaks
                if '\n' in cell.text:
                    cell.text = cell.text.replace('\n', '\n')
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to memory buffer
        doc_buffer = io.BytesIO()
        self.doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()